from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def write_formatted_docx(text, output_path):
    doc = Document()
    lines = text.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            doc.add_paragraph()
            i += 1
            continue

        # Big section headings: lines starting with ##
        if line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Dark Blue
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Horizontal rule
        if line.startswith('--'):
            doc.add_paragraph()
            i += 1
            continue

        # Bullet list (black color) with reduced spacing
        if line.startswith('- '):
            p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            # Default bullet text color black (no color set)
            i += 1
            continue

        # Bold label + value (on same line) e.g. "**Version:** 1.0.0"
        if line.startswith('**') and ':**' in line:
            split_index = line.find(':**')
            label = line[2:split_index].strip()
            value = line[split_index+3:].strip()

            p = doc.add_paragraph()
            run_label = p.add_run(label + ': ')
            run_label.bold = True
            run_label.font.size = Pt(12)
            run_label.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)  # Purple

            run_value = p.add_run(value)
            run_value.font.size = Pt(12)
            run_value.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Bold standalone line (like **Functional**) without colon
        if line.startswith('**') and line.endswith('**') and ':**' not in line:
            heading_text = line.strip('*').strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.italic = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)  # Medium Blue
            p.paragraph_format.space_after = Pt(3)
            i += 1

            # Indented paragraphs below heading
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line or next_line.startswith('## ') or next_line.startswith('**'):
                    break

                p2 = doc.add_paragraph(next_line)
                p2.paragraph_format.left_indent = Inches(0.25)
                p2.paragraph_format.space_after = Pt(2)
                p2.style = 'Normal'
                p2.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black text
                i += 1
            continue

        # Default paragraph
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(3)
        i += 1

    doc.save(output_path)
